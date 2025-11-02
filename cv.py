import io
import re
import json
import logging
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass, asdict
from io import BytesIO

import fitz  # PyMuPDF
import docx
import pytesseract
from PIL import Image
import pdfplumber
from sentence_transformers import SentenceTransformer
from keybert import KeyBERT

logger = logging.getLogger(__name__)

@dataclass
class CVProfile:
    """Structured CV profile data"""
    name: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    location: Optional[str] = None
    summary: Optional[str] = None
    skills: List[str] = None
    experience: List[Dict[str, Any]] = None
    education: List[Dict[str, Any]] = None
    certifications: List[str] = None
    keywords: List[str] = None
    raw_text: str = ""
    confidence_score: float = 0.0
    
    def __post_init__(self):
        if self.skills is None:
            self.skills = []
        if self.experience is None:
            self.experience = []
        if self.education is None:
            self.education = []
        if self.certifications is None:
            self.certifications = []
        if self.keywords is None:
            self.keywords = []

class CVProcessor:
    """Enhanced CV processing with structured data extraction"""
    
    def __init__(self, 
                 model_path: str = "all-MiniLM-L6-v2",
                 tesseract_path: Optional[str] = None):
        self.model = SentenceTransformer(model_path)
        self.kw_model = KeyBERT(model_path)
        
        if tesseract_path:
            pytesseract.pytesseract.tesseract_cmd = tesseract_path
        
        # Skill categories for better extraction
        self.skill_patterns = {
            'programming': r'\b(Python|Java|JavaScript|TypeScript|C\+\+|C#|Go|Rust|Ruby|PHP|Swift|Kotlin|Scala|R|MATLAB|SAS)\b',
            'web_frameworks': r'\b(React|Angular|Vue\.?js|Node\.?js|Django|Flask|FastAPI|Spring|Express|Next\.?js|Nuxt\.?js|Laravel|Rails)\b',
            'databases': r'\b(MySQL|PostgreSQL|MongoDB|Redis|Cassandra|DynamoDB|Oracle|SQL Server|SQLite|Elasticsearch|Neo4j)\b',
            'cloud_aws': r'\b(AWS|EC2|S3|Lambda|RDS|DynamoDB|CloudFormation|ECS|EKS|SageMaker)\b',
            'cloud_general': r'\b(Azure|GCP|Google Cloud|Kubernetes|Docker|Terraform|Ansible|Jenkins|GitLab CI|GitHub Actions)\b',
            'ml_ai': r'\b(TensorFlow|PyTorch|Scikit-learn|Pandas|NumPy|Matplotlib|Seaborn|Jupyter|Keras|OpenCV|NLTK|SpaCy|Hugging Face)\b',
            'tools': r'\b(Git|GitHub|GitLab|Jira|Confluence|Slack|Figma|Adobe|Photoshop|VS Code|IntelliJ|Eclipse)\b'
        }
        
        self.contact_patterns = {
            'email': r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
            'phone': r'(?:\+?1[-.\s]?)?(?:\(?[0-9]{3}\)?[-.\s]?)?[0-9]{3}[-.\s]?[0-9]{4}',
            'linkedin': r'(?:linkedin\.com/in/|linkedin\.com/pub/)([a-zA-Z0-9\-]+)',
            'github': r'(?:github\.com/)([a-zA-Z0-9\-]+)'
        }

    async def process_cv(self, file_content: bytes, filename: str) -> CVProfile:
        """Main CV processing pipeline"""
        try:
            # Extract text and tables
            raw_text, tables = await self._extract_content(file_content, filename)
            
            if not raw_text.strip():
                raise ValueError("No text content extracted from CV")
            
            # Parse structured data
            profile = await self._parse_cv_structure(raw_text, tables)
            profile.raw_text = raw_text
            
            # Extract keywords using KeyBERT
            profile.keywords = await self._extract_keywords(raw_text, tables)
            
            # Calculate confidence score
            profile.confidence_score = self._calculate_confidence(profile)
            
            logger.info(f"Successfully processed CV: {len(profile.skills)} skills, "
                       f"{len(profile.experience)} experience entries, "
                       f"confidence: {profile.confidence_score:.2f}")
            
            return profile
            
        except Exception as e:
            logger.error(f"CV processing failed: {e}")
            raise

    async def _extract_content(self, file_content: bytes, filename: str) -> Tuple[str, List[Dict]]:
        """Extract text and tables from various file formats"""
        file_ext = filename.lower().split('.')[-1]
        
        if file_ext == 'pdf':
            return await self._extract_from_pdf(file_content)
        elif file_ext in ['doc', 'docx']:
            return await self._extract_from_docx(file_content)
        elif file_ext in ['png', 'jpg', 'jpeg']:
            return await self._extract_from_image(file_content)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")

    async def _extract_from_pdf(self, file_content: bytes) -> Tuple[str, List[Dict]]:
        """Enhanced PDF extraction with table support"""
        text_output = ""
        tables = []
        
        try:
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    # Extract text
                    page_text = page.extract_text()
                    if page_text:
                        text_output += f"\n--- Page {page_num + 1} ---\n{page_text}"
                    
                    # Extract tables
                    page_tables = page.extract_tables()
                    for table in page_tables:
                        if table:
                            table_data = []
                            for row in table:
                                if row:
                                    clean_row = [cell.strip() if cell else "" for cell in row]
                                    if any(clean_row):
                                        table_data.append(clean_row)
                            
                            if table_data:
                                tables.append({
                                    'page': page_num + 1,
                                    'data': table_data,
                                    'text': '\n'.join([' | '.join(row) for row in table_data])
                                })
                    
                    # Fallback to OCR if no text extracted
                    if not page_text or len(page_text.strip()) < 50:
                        try:
                            pil_image = page.to_image(resolution=300).original
                            ocr_text = pytesseract.image_to_string(pil_image)
                            if ocr_text.strip():
                                text_output += f"\n--- OCR Page {page_num + 1} ---\n{ocr_text}"
                        except Exception as ocr_error:
                            logger.warning(f"OCR failed for page {page_num + 1}: {ocr_error}")
            
            return text_output, tables
            
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            # Fallback to PyMuPDF
            try:
                doc = fitz.open(stream=file_content, filetype="pdf")
                text_output = ""
                for page_num in range(doc.page_count):
                    page = doc[page_num]
                    text_output += f"\n--- Page {page_num + 1} ---\n{page.get_text()}"
                doc.close()
                return text_output, []
            except Exception as fallback_error:
                raise Exception(f"All PDF extraction methods failed: {e}, {fallback_error}")

    async def _extract_from_docx(self, file_content: bytes) -> Tuple[str, List[Dict]]:
        """Extract from Word documents"""
        try:
            doc = docx.Document(BytesIO(file_content))
            text_parts = []
            tables = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())
            
            # Extract tables
            for table_num, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_data:
                        table_data.append(row_data)
                
                if table_data:
                    tables.append({
                        'table_num': table_num + 1,
                        'data': table_data,
                        'text': '\n'.join([' | '.join(row) for row in table_data])
                    })
            
            return '\n'.join(text_parts), tables
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise

    async def _extract_from_image(self, file_content: bytes) -> Tuple[str, List[Dict]]:
        """Extract from image files using OCR"""
        try:
            image = Image.open(io.BytesIO(file_content))
            text = pytesseract.image_to_string(image, config='--psm 6')
            return text, []
        except Exception as e:
            logger.error(f"Image OCR failed: {e}")
            raise

    async def _parse_cv_structure(self, text: str, tables: List[Dict]) -> CVProfile:
        """Parse structured information from CV text"""
        profile = CVProfile()
        
        # Extract contact information
        profile.email = self._extract_contact_info(text, 'email')
        profile.phone = self._extract_contact_info(text, 'phone')
        profile.name = self._extract_name(text)
        profile.location = self._extract_location(text)
        
        # Extract skills
        profile.skills = self._extract_skills(text, tables)
        
        # Extract experience
        profile.experience = self._extract_experience(text)
        
        # Extract education
        profile.education = self._extract_education(text)
        
        # Extract certifications
        profile.certifications = self._extract_certifications(text)
        
        # Extract summary
        profile.summary = self._extract_summary(text)
        
        return profile

    def _extract_contact_info(self, text: str, info_type: str) -> Optional[str]:
        """Extract specific contact information"""
        if info_type not in self.contact_patterns:
            return None
        
        pattern = self.contact_patterns[info_type]
        matches = re.findall(pattern, text, re.IGNORECASE)
        return matches[0] if matches else None

    def _extract_name(self, text: str) -> Optional[str]:
        """Extract name from CV (heuristic approach)"""
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        if not lines:
            return None
        
        # Usually the first meaningful line is the name
        first_line = lines[0]
        
        # Filter out obvious non-names
        if any(keyword in first_line.lower() for keyword in 
               ['resume', 'cv', 'curriculum', 'page', 'confidential']):
            return None
        
        # Simple name validation (2-4 words, mostly alphabetic)
        words = first_line.split()
        if 2 <= len(words) <= 4 and all(len(w) > 1 and w.replace('-', '').replace("'", '').isalpha() for w in words):
            return first_line
        
        return None

    def _extract_location(self, text: str) -> Optional[str]:
        """Extract location information"""
        location_patterns = [
            r'\b([A-Z][a-z]+,\s*[A-Z]{2})\b',  # City, State
            r'\b([A-Z][a-z]+,\s*[A-Z][a-z]+)\b',  # City, Country
            r'\b(\d{5})\b'  # ZIP code
        ]
        
        for pattern in location_patterns:
            matches = re.findall(pattern, text)
            if matches:
                return matches[0]
        
        return None

    def _extract_skills(self, text: str, tables: List[Dict]) -> List[str]:
        """Extract technical skills using patterns and tables"""
        skills = set()
        
        # Combine text and table content
        combined_text = text
        for table in tables:
            combined_text += "\n" + table.get('text', '')
        
        # Extract using predefined patterns
        for category, pattern in self.skill_patterns.items():
            matches = re.findall(pattern, combined_text, re.IGNORECASE)
            skills.update([match.strip() for match in matches])
        
        # Look for skills sections
        skills_section = self._extract_section(text, ['skills', 'technical skills', 'technologies'])
        if skills_section:
            # Common skill separators
            skill_items = re.split(r'[,;•·\n\|]', skills_section)
            for item in skill_items:
                clean_item = item.strip()
                if 2 <= len(clean_item) <= 25 and not clean_item.lower() in ['and', 'or', 'with']:
                    skills.add(clean_item)
        
        return sorted(list(skills))

    def _extract_experience(self, text: str) -> List[Dict[str, Any]]:
        """Extract work experience entries"""
        experience = []
        
        # Look for experience section
        exp_section = self._extract_section(text, ['experience', 'work experience', 'employment', 'professional experience'])
        
        if not exp_section:
            return experience
        
        # Split by common job entry patterns
        job_patterns = [
            r'\n(?=\w+.*(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec))',
            r'\n(?=\d{4}\s*[-–]\s*(?:\d{4}|Present))',
            r'\n(?=[A-Z][a-z]+\s+\w+.*\d{4})'
        ]
        
        entries = [exp_section]
        for pattern in job_patterns:
            new_entries = []
            for entry in entries:
                splits = re.split(pattern, entry)
                new_entries.extend(splits)
            entries = new_entries
        
        for entry in entries:
            if len(entry.strip()) > 50:  # Meaningful entry
                exp_data = self._parse_experience_entry(entry.strip())
                if exp_data:
                    experience.append(exp_data)
        
        return experience[:5]  # Limit to recent 5 positions

    def _parse_experience_entry(self, entry: str) -> Optional[Dict[str, Any]]:
        """Parse individual experience entry"""
        lines = [line.strip() for line in entry.split('\n') if line.strip()]
        if len(lines) < 2:
            return None
        
        # Extract dates
        date_pattern = r'(\d{4})\s*[-–]\s*(\d{4}|Present|Current)'
        date_match = re.search(date_pattern, entry)
        
        exp_data = {
            'title': None,
            'company': None,
            'duration': None,
            'start_date': None,
            'end_date': None,
            'description': None
        }
        
        if date_match:
            exp_data['start_date'] = date_match.group(1)
            exp_data['end_date'] = date_match.group(2)
            exp_data['duration'] = date_match.group(0)
        
        # First line often contains title and/or company
        first_line = lines[0]
        exp_data['title'] = first_line.split(' at ')[0] if ' at ' in first_line else first_line
        
        if ' at ' in first_line:
            exp_data['company'] = first_line.split(' at ')[-1]
        
        # Description from remaining lines
        if len(lines) > 1:
            exp_data['description'] = '\n'.join(lines[1:])
        
        return exp_data

    def _extract_education(self, text: str) -> List[Dict[str, Any]]:
        """Extract education information"""
        education = []
        
        edu_section = self._extract_section(text, ['education', 'academic background', 'qualifications'])
        if not edu_section:
            return education
        
        # Common degree patterns
        degree_patterns = [
            r'\b(Bachelor|Master|PhD|B\.?S\.?|M\.?S\.?|B\.?A\.?|M\.?A\.?|MBA|B\.?Tech|M\.?Tech)\b',
            r'\b(Associate|Diploma|Certificate)\b'
        ]
        
        lines = edu_section.split('\n')
        current_entry = {}
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if line contains degree
            for pattern in degree_patterns:
                if re.search(pattern, line, re.IGNORECASE):
                    if current_entry:
                        education.append(current_entry)
                    current_entry = {'degree': line, 'institution': None, 'year': None}
                    
                    # Extract year
                    year_match = re.search(r'\b(19|20)\d{2}\b', line)
                    if year_match:
                        current_entry['year'] = year_match.group(0)
                    break
            else:
                # This line might be institution or additional info
                if current_entry and not current_entry.get('institution'):
                    current_entry['institution'] = line
        
        if current_entry:
            education.append(current_entry)
        
        return education

    def _extract_certifications(self, text: str) -> List[str]:
        """Extract certifications"""
        cert_section = self._extract_section(text, ['certifications', 'certificates', 'licenses'])
        if not cert_section:
            return []
        
        # Split by common separators
        cert_items = re.split(r'[,;\n•·]', cert_section)
        certifications = []
        
        for item in cert_items:
            clean_item = item.strip()
            if 3 <= len(clean_item) <= 100:
                certifications.append(clean_item)
        
        return certifications

    def _extract_summary(self, text: str) -> Optional[str]:
        """Extract professional summary"""
        summary_keywords = ['summary', 'objective', 'profile', 'about', 'overview']
        
        for keyword in summary_keywords:
            summary = self._extract_section(text, [keyword])
            if summary and len(summary) > 50:
                return summary[:500]  # Limit length
        
        return None

    def _extract_section(self, text: str, section_names: List[str]) -> Optional[str]:
        """Extract specific section from CV text"""
        for section_name in section_names:
            # Pattern to match section headers
            pattern = rf'\n\s*{re.escape(section_name)}\s*:?\s*\n(.*?)(?=\n\s*[A-Z][A-Z\s]*:|\Z)'
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            
            if match:
                return match.group(1).strip()
            
            # Alternative pattern
            pattern = rf'\b{re.escape(section_name)}\b\s*:?\s*(.*?)(?=\n\s*[A-Z][A-Z\s]*:|\Z)'
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            
            if match:
                section_text = match.group(1).strip()
                if len(section_text) > 20:
                    return section_text
        
        return None

    async def _extract_keywords(self, text: str, tables: List[Dict], top_n: int = 30) -> List[str]:
        """Extract keywords using KeyBERT"""
        try:
            # Combine text and table content
            combined_text = text
            for table in tables:
                combined_text += "\n" + table.get('text', '')
            
            # Extract keywords
            keywords_with_scores = self.kw_model.extract_keywords(
                combined_text,
                keyphrase_ngram_range=(1, 2),
                stop_words='english',
                top_k=top_n,
                use_mmr=True,
                diversity=0.5
            )
            
            return [kw for kw, score in keywords_with_scores]
            
        except Exception as e:
            logger.error(f"Keyword extraction failed: {e}")
            return []

    def _calculate_confidence(self, profile: CVProfile) -> float:
        """Calculate confidence score for extracted data"""
        score = 0.0
        
        # Contact information (20%)
        if profile.email:
            score += 0.1
        if profile.phone:
            score += 0.05
        if profile.name:
            score += 0.05
        
        # Skills (30%)
        skill_score = min(len(profile.skills) / 10, 1.0) * 0.3
        score += skill_score
        
        # Experience (25%)
        exp_score = min(len(profile.experience) / 3, 1.0) * 0.25
        score += exp_score
        
        # Education (15%)
        edu_score = min(len(profile.education) / 2, 1.0) * 0.15
        score += edu_score
        
        # Content quality (10%)
        if len(profile.raw_text) > 500:
            score += 0.05
        if profile.summary:
            score += 0.05
        
        return round(score, 2)

    def get_job_search_text(self, profile: CVProfile) -> str:
        """Generate optimized text for job searching"""
        search_components = []
        
        # High-priority skills
        if profile.skills:
            search_components.append(" ".join(profile.skills[:15]))
        
        # Professional summary keywords
        if profile.summary:
            # Extract key phrases from summary
            summary_keywords = re.findall(r'\b[A-Za-z]{3,}\b', profile.summary)
            search_components.append(" ".join(summary_keywords[:10]))
        
        # Recent job titles
        if profile.experience:
            recent_titles = [exp.get('title', '') for exp in profile.experience[:2] if exp.get('title')]
            search_components.extend(recent_titles)
        
        # Top keywords
        if profile.keywords:
            search_components.append(" ".join(profile.keywords[:10]))
        
        return " ".join(search_components)

    def to_dict(self, profile: CVProfile) -> Dict[str, Any]:
        """Convert profile to dictionary"""
        return asdict(profile)